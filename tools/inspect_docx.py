from pathlib import Path
import sys
from docx import Document


path = (
    Path(sys.argv[1]).resolve()
    if len(sys.argv) > 1
    else Path(__file__).resolve().parents[1] / "Relatorio_FootCentral.docx"
)
doc = Document(path)

for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    drawings = paragraph._p.xpath('.//w:drawing')
    if text or drawings:
        marker = f" <DRAWING x{len(drawings)}>" if drawings else ""
        print(f"P{index:03d} [{paragraph.style.name}] {text}{marker}")

for table_index, table in enumerate(doc.tables):
    table_drawings = table._tbl.xpath('.//w:drawing')
    drawing_marker = f" <DRAWING x{len(table_drawings)}>" if table_drawings else ""
    print(f"TABLE {table_index}{drawing_marker}")
    for row_index, row in enumerate(table.rows):
        print(f"  R{row_index}: " + " | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
