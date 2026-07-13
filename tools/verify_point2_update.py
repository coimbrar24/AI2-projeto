from pathlib import Path
from zipfile import ZipFile

from docx import Document


root = Path(__file__).resolve().parents[1]
source_path = root / "Relatorio_FootCentral.docx"
output_path = root / "Relatorio_FootCentral_atualizado.docx"

source = Document(source_path)
output = Document(output_path)

assert len(source.paragraphs) == len(output.paragraphs)
assert len(source.tables) == len(output.tables)

for index in list(range(0, 74)) + list(range(128, len(source.paragraphs))):
    assert source.paragraphs[index].text == output.paragraphs[index].text, index
    assert source.paragraphs[index].style.name == output.paragraphs[index].style.name, index

assert [[cell.text for cell in row.cells] for row in source.tables[0].rows] == [
    [cell.text for cell in row.cells] for row in output.tables[0].rows
]

with ZipFile(output_path) as archive:
    bad_file = archive.testzip()
    assert bad_file is None, bad_file
    media = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert media, "The original embedded screenshot was not preserved"

assert output.paragraphs[74].text.startswith("2. Descrição")
assert output.paragraphs[128].text.startswith("3. Conclusões")
assert len(output.tables[3].rows) == 12

print("Structural QA passed")
print(f"Paragraphs: {len(output.paragraphs)}")
print(f"Tables: {len(output.tables)}")
print(f"Embedded media: {len(media)}")
