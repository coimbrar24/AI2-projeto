from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Relatorio_FootCentral.docx"
OUTPUT = ROOT / "Relatorio_FootCentral_atualizado.docx"


def replace_paragraph_text(paragraph, text):
    """Replace text while preserving paragraph properties and representative run formatting."""
    run_properties = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            run_properties = deepcopy(run._r.rPr)
            break

    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)

    run = OxmlElement("w:r")
    if run_properties is not None:
        run.append(run_properties)
    text_element = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run.append(text_element)
    paragraph._p.append(run)


def replace_cell_text(cell, text, formatting_source=None):
    paragraph = cell.paragraphs[0]
    source = formatting_source or paragraph
    run_properties = None
    for run in source.runs:
        if run._r.rPr is not None:
            run_properties = deepcopy(run._r.rPr)
            break

    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)

    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


doc = Document(SOURCE)

replacements = {
    76: (
        "Numa primeira etapa foi definida uma arquitetura cliente-servidor. O frontend, desenvolvido em React e Vite, "
        "comunica por HTTP com uma API REST criada em Node.js e Express. O backend centraliza o acesso ao PostgreSQL, "
        "através do Sequelize, e à API Football-Data.org, devolvendo respostas em formato JSON."
    ),
    78: (
        "O frontend foi organizado em páginas, componentes reutilizáveis, contextos e um serviço Axios comum. O backend "
        "foi dividido em rotas, controladores, serviços, modelos, middleware e configuração. Esta separação mantém a "
        "interface independente da persistência e da fonte externa e facilita a manutenção do código."
    ),
    81: (
        "O backend foi desenvolvido com Express e disponibiliza uma API REST sob o prefixo /api. Foi também criado o "
        "endpoint /health, que confirma o estado do serviço. A aplicação aceita JSON, utiliza CORS e possui middleware "
        "comum para erros internos e rotas não encontradas."
    ),
    82: "As rotas atualmente implementadas são:",
    84: (
        "Os controladores validam parâmetros como datas, limites, termos de pesquisa e identificadores, delegando o acesso "
        "a dados aos serviços. O serviço de futebol normaliza as respostas, limita resultados e mantém uma cache em memória "
        "com deduplicação de pedidos simultâneos. O serviço de equipas suporta pesquisa textual e consulta do detalhe de cada clube."
    ),
    85: (
        "A configuração é lida a partir de variáveis de ambiente: porta, ligação ao PostgreSQL, segredo e validade dos JWT, "
        "URL e chave da Football-Data.org. Na base de dados existem os modelos User e Favorite; cada favorito pertence a um "
        "utilizador e pode representar uma equipa ou um jogo, com restrição de duplicados por utilizador, tipo e identificador externo."
    ),
    87: (
        "O frontend foi criado em React e organizado em páginas e componentes reutilizáveis. A página principal apresenta uma "
        "barra lateral com competições, uma área central com jogos agrupados por competição e um painel lateral com notícias e favoritos."
    ),
    88: (
        "Na página principal, o utilizador pode consultar os jogos de ontem, hoje ou amanhã, filtrar por competição e ativar a "
        "visualização de jogos ao vivo. Cada mudança de data origina um pedido a /api/matches; os resultados, até ao limite de 50, "
        "são agrupados no cliente e apresentados com estado, hora ou resultado, equipas e respetivos emblemas."
    ),
    96: (
        "Foram implementados estados distintos de carregamento, erro e ausência de resultados. A interface adapta também a mensagem "
        "ao filtro selecionado e informa o utilizador quando é necessário iniciar sessão para utilizar funcionalidades reservadas."
    ),
    97: (
        "A barra de pesquisa encaminha para uma página global de pesquisa de equipas. A consulta exige pelo menos dois caracteres e "
        "permite encontrar clubes pelo nome, abreviatura ou país. Ao selecionar um resultado, é apresentada uma página com emblema, "
        "fundação, estádio, treinador, competições em curso e plantel agrupado por posição."
    ),
    99: "Figura 5 - Filtragem de jogos por data, competição e estado ao vivo.",
    100: "Figura 6 - Pesquisa global de equipas.",
    101: "Figura 7 - Informação geral e plantel de uma equipa.",
    103: (
        "Para suportar contas de utilizador foi criado o modelo User, associado à tabela users, e o modelo Favorite, associado à "
        "tabela favorites. A relação permite guardar equipas e jogos por utilizador, ficando as rotas de favoritos protegidas por autenticação."
    ),
    104: "Figura 3 - Modelo de dados do utilizador e dos favoritos.",
    105: (
        "Durante o registo, o backend valida nome, email e palavra-passe, exige um mínimo de seis caracteres e impede emails duplicados. "
        "A palavra-passe é processada com bcrypt antes de ser guardada, pelo que o valor original não é armazenado."
    ),
    106: (
        "No início de sessão, o sistema compara a palavra-passe recebida com o hash existente. Quando as credenciais são válidas, "
        "é emitido um JWT com validade configurável. O endpoint /api/auth/me utiliza middleware de autenticação para devolver os dados do utilizador."
    ),
    110: (
        "No frontend, o AuthContext mantém o utilizador e o token no localStorage, e um interceptor do Axios acrescenta o cabeçalho "
        "Authorization aos pedidos autenticados. A rota /dashboard está protegida e redireciona para /login sem sessão válida. Foi também "
        "criado um FavoritesContext que comunica com a API, embora ainda falte ligá-lo ao Provider principal, aos botões e à página de favoritos."
    ),
    113: (
        "A fonte externa efetivamente integrada nesta versão é a Football-Data.org. O backend utiliza o cabeçalho X-Auth-Token para "
        "consultar jogos e equipas, mantendo a chave fora do frontend. Quando a chave não está configurada ou a API não responde, "
        "o serviço devolve um erro controlado."
    ),
    114: (
        "A pesquisa de equipas consulta até duas páginas de 500 registos, normaliza acentos e maiúsculas e ordena os resultados, dando "
        "prioridade aos nomes iniciados pelo termo procurado. A lista de equipas é mantida em cache durante uma hora; o detalhe de cada "
        "equipa, incluindo época, treinador, competições e plantel, durante dez minutos."
    ),
    115: (
        "Os pedidos de jogos utilizam uma cache em memória de um minuto e reutilizam pedidos iguais que estejam em curso. O painel de "
        "notícias contém atualmente conteúdo estático de demonstração; NewsAPI, TheSportsDB e o feed RSS referidos no planeamento ainda "
        "não estão ligados ao código desta versão."
    ),
    117: (
        "Em 13 de julho de 2026 foi executada a compilação de produção do frontend através do Vite. O processo terminou com sucesso, "
        "transformou 99 módulos e gerou os ficheiros finais da aplicação sem erros."
    ),
    118: (
        "A análise estática com Oxlint terminou sem erros e apresentou seis advertências: uma importação não utilizada no Dashboard, "
        "duas advertências relativas à exportação conjunta de componentes e funções nos contextos e três relacionadas com dependências "
        "do useMemo no FavoritesContext. Foi ainda verificada, sem erros, a sintaxe dos 23 ficheiros JavaScript do backend."
    ),
    119: (
        "Ainda não existe uma suite de testes automatizados. A validação integral do backend requer PostgreSQL em execução, uma chave "
        "válida da Football-Data.org e testes dos fluxos autenticados. Por esse motivo, a compilação e a análise estática confirmam a "
        "consistência do código, mas não substituem testes de integração e testes end-to-end."
    ),
    120: "Durante a revisão do estado atual foram identificados os seguintes pontos a concluir ou melhorar:",
    121: "integrar o FavoritesProvider na aplicação e ligar os botões de favorito ao FavoritesContext;",
    122: "concluir a página de favoritos e apresentar as equipas e os jogos guardados;",
    123: "substituir as notícias estáticas por uma fonte de dados externa;",
    124: "concluir as páginas de competições e notícias e evoluir a área reservada do utilizador;",
    125: "corrigir os avisos do Oxlint e os textos com problemas de codificação de caracteres;",
    126: "limitar as origens autorizadas pela configuração de CORS;",
    127: "adicionar testes unitários, de integração e end-to-end para os principais fluxos.",
}

for paragraph_index, text in replacements.items():
    replace_paragraph_text(doc.paragraphs[paragraph_index], text)

replace_cell_text(
    doc.tables[1].cell(0, 0),
    "flowchart LR\n"
    "    U[Utilizador] --> F[Frontend React e Vite]\n"
    "    F -->|HTTP e JSON| B[API Node.js e Express]\n"
    "    B --> D[(PostgreSQL: utilizadores e favoritos)]\n"
    "    B --> FD[Football-Data.org]\n"
    "    B --> C[Cache em memória]",
)

replace_cell_text(
    doc.tables[6].cell(0, 0),
    "User\n"
    "---------------------------------\n"
    "id          INTEGER, chave primária\n"
    "name        VARCHAR(100), obrigatório\n"
    "email       VARCHAR(150), único\n"
    "password    VARCHAR, hash bcrypt\n"
    "createdAt   data e hora\n"
    "updatedAt   data e hora\n"
    "             1 ----- N\n"
    "Favorite\n"
    "---------------------------------\n"
    "id          INTEGER, chave primária\n"
    "userId      INTEGER, chave estrangeira\n"
    "type        VARCHAR(10), team ou match\n"
    "externalId  INTEGER, identificador externo\n"
    "data        JSONB, resumo do favorito\n"
    "createdAt   data e hora\n"
    "updatedAt   data e hora\n"
    "UNIQUE (userId, type, externalId)",
)

endpoint_rows = [
    ("GET", "/health", "Confirmar o estado do serviço"),
    ("POST", "/api/auth/register", "Registar um utilizador"),
    ("POST", "/api/auth/login", "Autenticar um utilizador"),
    ("GET", "/api/auth/me", "Obter o utilizador autenticado"),
    ("GET", "/api/matches?date=&limit=", "Consultar jogos de uma data"),
    ("GET", "/api/matches/results?dateFrom=&dateTo=", "Consultar resultados num intervalo até 10 dias"),
    ("GET", "/api/teams/search?q=", "Pesquisar equipas"),
    ("GET", "/api/teams/:id", "Consultar detalhe, competições e plantel"),
    ("GET", "/api/favorites", "Listar os favoritos do utilizador"),
    ("POST", "/api/favorites", "Guardar uma equipa ou um jogo"),
    ("DELETE", "/api/favorites/:type/:externalId", "Remover um favorito"),
]

table = doc.tables[3]
header_sources = [cell.paragraphs[0] for cell in table.rows[0].cells]
body_sources = [cell.paragraphs[0] for cell in table.rows[1].cells]

while len(table.rows) - 1 < len(endpoint_rows):
    table.add_row()
while len(table.rows) - 1 > len(endpoint_rows):
    table._tbl.remove(table.rows[-1]._tr)

for row, values in zip(table.rows[1:], endpoint_rows):
    for column_index, (cell, value) in enumerate(zip(row.cells, values)):
        replace_cell_text(cell, value, body_sources[column_index])

doc.save(OUTPUT)
print(OUTPUT)
